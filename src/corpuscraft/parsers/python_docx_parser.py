from __future__ import annotations

import logging
import re
from pathlib import Path

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from corpuscraft.config import ParserConfig
from corpuscraft.models import ParsedDocument
from corpuscraft.parsers.base import BaseParser

logger = logging.getLogger(__name__)

_HEADING_RE = re.compile(r"^Heading\s+(\d+)$", re.IGNORECASE)


def _iter_block_items(document):
    """Yield Paragraph and Table objects in document order."""
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def _heading_level(style_name: str | None) -> int | None:
    if not style_name:
        return None
    if style_name.lower() == "title":
        return 1
    m = _HEADING_RE.match(style_name)
    return int(m.group(1)) if m else None


def _paragraph_to_markdown(para: Paragraph) -> str:
    text = para.text.strip()
    if not text:
        return ""

    level = _heading_level(para.style.name if para.style else None)
    if level is not None:
        return f"{'#' * min(level, 6)} {text}"

    style_name = (para.style.name if para.style else "").lower()
    if "list" in style_name:
        return f"- {text}"

    return text


def _table_to_markdown(table: Table) -> str:
    if not table.rows:
        return ""

    rows: list[list[str]] = []
    for row in table.rows:
        cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
        rows.append(cells)

    width = max(len(r) for r in rows)
    rows = [r + [""] * (width - len(r)) for r in rows]

    lines = ["| " + " | ".join(rows[0]) + " |"]
    lines.append("| " + " | ".join(["---"] * width) + " |")
    for row in rows[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


class PythonDocxParser(BaseParser):
    """Lightweight DOCX parser using python-docx.

    Walks the document body in order, converting Heading styles to markdown
    headers, tables to markdown tables, and List styles to bullet items.
    """

    def __init__(self, config: ParserConfig):
        self.config = config

    def parse_file(self, path: Path) -> ParsedDocument:
        logger.info(f"Parsing {path.name} with python-docx")

        doc = Document(str(path))

        parts: list[str] = []
        n_paragraphs = 0
        n_tables = 0

        for block in _iter_block_items(doc):
            if isinstance(block, Paragraph):
                md = _paragraph_to_markdown(block)
                if md:
                    parts.append(md)
                    n_paragraphs += 1
            elif isinstance(block, Table):
                md = _table_to_markdown(block)
                if md:
                    parts.append(md)
                    n_tables += 1

        # Extract content from headers/footers in each section
        for section in doc.sections:
            for header_para in section.header.paragraphs:
                text = header_para.text.strip()
                if text:
                    parts.insert(0, f"<!-- header: {text} -->")
            for footer_para in section.footer.paragraphs:
                text = footer_para.text.strip()
                if text:
                    parts.append(f"<!-- footer: {text} -->")

        content = "\n\n".join(parts)

        core = doc.core_properties
        metadata = {
            "title": core.title or "",
            "author": core.author or "",
            "created": core.created.isoformat() if core.created else "",
            "modified": core.modified.isoformat() if core.modified else "",
            "paragraph_count": n_paragraphs,
            "table_count": n_tables,
            "section_count": len(doc.sections),
            "file_size": path.stat().st_size,
        }

        return ParsedDocument(
            content=content,
            source_path=path,
            pipeline="python_docx",
            metadata=metadata,
        )
